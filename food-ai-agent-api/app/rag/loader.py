"""Document loader - extracts text from PDF, DOCX, Markdown, TXT files."""
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class RawDocument:
    content: str
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """Load documents from various formats into plain text."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}

    async def load(self, file_path: str, **extra_metadata) -> list[RawDocument]:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            docs = self._load_pdf(path)
        elif ext == ".docx":
            docs = self._load_docx(path)
        elif ext in (".md", ".txt"):
            docs = self._load_text(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        # 한국어 정규화 + 메타데이터 보강
        for doc in docs:
            doc.content = unicodedata.normalize("NFKC", doc.content)
            doc.content = self._clean_text(doc.content)
            doc.metadata.update({"source_file": path.name, **extra_metadata})

        return docs

    def _load_pdf(self, path: Path) -> list[RawDocument]:
        docs = []
        with fitz.open(str(path)) as pdf:
            title = pdf.metadata.get("title", path.stem) if pdf.metadata else path.stem
            for page_num, page in enumerate(pdf):
                text = page.get_text("text").strip()
                if text:
                    docs.append(RawDocument(
                        content=text,
                        metadata={"title": title, "page": page_num + 1},
                    ))
        # 단일 문서로 합칠 수도 있지만, 페이지별로 유지하여 출처 추적 용이
        if docs:
            merged = RawDocument(
                content="\n\n".join(d.content for d in docs),
                metadata={"title": title, "total_pages": len(docs)},
            )
            return [merged]
        return []

    def _load_docx(self, path: Path) -> list[RawDocument]:
        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return []
        return [RawDocument(
            content="\n\n".join(paragraphs),
            metadata={"title": path.stem},
        )]

    def _load_text(self, path: Path) -> list[RawDocument]:
        content = path.read_text(encoding="utf-8").strip()
        if not content:
            return []
        return [RawDocument(
            content=content,
            metadata={"title": path.stem},
        )]

    @staticmethod
    def _clean_text(text: str) -> str:
        """Remove excessive whitespace while preserving paragraph breaks."""
        import re
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
